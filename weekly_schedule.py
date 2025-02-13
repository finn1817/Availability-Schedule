import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageDraw, ImageFont
from docx import Document


class WeeklyScheduleGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Weekly Schedule Generator")
        self.root.geometry("400x250")

        self.load_button = tk.Button(root, text="Load Excel Schedule", command=self.load_schedule)
        self.load_button.pack(pady=10)

        self.generate_availability_button = tk.Button(root, text="Generate Availability", command=self.generate_availability)
        self.generate_availability_button.pack(pady=10)

        self.generate_schedule_button = tk.Button(root, text="Generate Schedule", command=self.generate_schedule)
        self.generate_schedule_button.pack(pady=10)

        self.save_button = tk.Button(root, text="Save Word File", command=self.save_word_file)
        self.save_button.pack(pady=10)

        self.availability_data = None  # who is available for each time slot
        self.schedule_data = None  # final worker for each time slot
        self.image_path = "weekly_schedule.png"  # path to save the picture

    def load_schedule(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if not file_path:
            return

        try:
            # load data into a pandas dataframe
            self.schedule_df = pd.read_excel(file_path)

            # make sure columns are on the excel sheet
            required_columns = {'First Name', 'Last Name', 'Email', 'Days Available'}
            if not required_columns.issubset(self.schedule_df.columns):
                messagebox.showerror("Error", f"Excel file must contain the following columns: {', '.join(required_columns)}")
                return

            messagebox.showinfo("Success", "Schedule file loaded successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load schedule file: {e}")
            self.schedule_df = None

    def generate_availability(self):
        if self.schedule_df is None:
            messagebox.showerror("Error", "Please load a schedule file first.")
            return

        # shifts and times (can change whenever)
        shifts = {
            "Sunday": ["12 PM - 4 PM", "4 PM - 7 PM", "7 PM - 10 PM", "10 PM - 12 AM"],
            "Monday": ["2 PM - 5 PM", "5 PM - 8 PM", "8 PM - 12 AM"],
            "Tuesday": ["2 PM - 5 PM", "5 PM - 8 PM", "8 PM - 12 AM"],
            "Wednesday": ["2 PM - 6 PM", "6 PM - 9 PM", "9 PM - 12 AM"],
            "Thursday": ["2 PM - 4 PM", "4 PM - 8 PM", "8 PM - 12 AM"],
            "Friday": ["2 PM - 7 PM", "7 PM - 9 PM", "9 PM - 12 AM"],
            "Saturday": ["12 PM - 4 PM", "4 PM - 8 PM", "8 PM - 12 AM"]
        }

        # prep availability data
        availability = {day: {shift: [] for shift in shifts[day]} for day in shifts.keys()}

        for _, row in self.schedule_df.iterrows():
            for day in availability.keys():
                if day in row["Days Available"]:
                    for shift in availability[day]:
                        availability[day][shift].append(f"{row['First Name']} {row['Last Name']}")

        self.availability_data = availability

        self.create_calendar_image(availability, shifts, title="Who is Available")
        messagebox.showinfo("Success", "Availability generated! Image saved as 'weekly_schedule.png'.")

    def generate_schedule(self):
        if not self.availability_data:
            messagebox.showerror("Error", "Please generate availability first.")
            return

        # start the final schedule with the same structure as availability
        schedule = {day: {shift: None for shift in shifts} for day, shifts in self.availability_data.items()}

        # call on one person per shift
        for day, shifts in self.availability_data.items():
            assigned_workers = set()  # keep track of workers already assigned
            for shift, workers in shifts.items():
                # assign the first available worker who isn't already assigned
                for worker in workers:
                    if worker not in assigned_workers:
                        schedule[day][shift] = worker
                        assigned_workers.add(worker)
                        break

                # if no one is available, over assign a worker (from the top of the list) to fill the spots
                if not schedule[day][shift] and workers:
                    schedule[day][shift] = workers[0]  # reassign the first worker in the list (queue from bottom to top)

        self.schedule_data = schedule
        self.create_calendar_image(schedule, {day: list(shifts.keys()) for day, shifts in self.availability_data.items()}, title="Final Schedule")
        messagebox.showinfo("Success", "Final schedule generated! Image saved as 'weekly_schedule.png'.")

    def create_calendar_image(self, data, shifts, title="Weekly Schedule"):
        # picture settings
        width, height = 1200, 1000
        header_height = 100
        color_bg = (255, 255, 255)
        color_header = (70, 130, 180)
        color_text = (0, 0, 0)

        # make the picture
        img = Image.new("RGB", (width, height), color_bg)
        draw = ImageDraw.Draw(img)

        # add title
        font_title = ImageFont.truetype("arial.ttf", 40)
        draw.text((width // 2 - 200, 10), title, fill=color_header, font=font_title)

        # draw the schedule by day and shift
        font_body = ImageFont.truetype("arial.ttf", 20)
        row_height = (height - header_height) // len(data)

        for i, (day, shift_data) in enumerate(data.items()):
            y_start = header_height + i * row_height

            # draw out day header
            draw.text((10, y_start + 10), day, fill=color_header, font=font_body)

            # add shifts and workers
            y_shift = y_start + 40
            for shift, worker in shift_data.items():
                shift_text = f"{shift}: {worker if worker else 'No one assigned'}"
                draw.text((40, y_shift), shift_text, fill=color_text, font=font_body)
                y_shift += 30

        # save the picture
        img.save(self.image_path)

    def save_word_file(self):
        if not self.schedule_data:
            messagebox.showerror("Error", "Please generate the final schedule first.")
            return

        try:
            # make the word document
            doc = Document()
            doc.add_heading("Weekly Schedule", level=1)

            for day, shift_data in self.schedule_data.items():
                doc.add_heading(day, level=2)
                for shift, worker in shift_data.items():
                    doc.add_paragraph(f"{shift}: {worker if worker else 'No one assigned'}")

            # save word file
            file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if file_path:
                doc.save(file_path)
                messagebox.showinfo("Success", f"Word file saved successfully at {file_path}.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save Word file: {e}")


if __name__ == "__main__":
    root = tk.Tk()
    app = WeeklyScheduleGenerator(root)
    root.mainloop()
